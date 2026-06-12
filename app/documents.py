"""文字起こし結果から Word(.docx) / Excel(.xlsx) を生成する（いずれもTC入り）。"""
from __future__ import annotations

import io

from . import formats

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


# ---------------------------------------------------------------- Word
def _set_japanese_font(doc, name: str = "Yu Mincho") -> None:
    from docx.oxml.ns import qn

    style = doc.styles["Normal"]
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)


def build_docx(
    paragraphs: list[dict], title: str = "文字起こし", subtitle: str = "", show_tc: bool = True
) -> bytes:
    """ブロックごとに本文を並べた Word 文書を bytes で返す。

    show_tc=True なら各ブロック先頭に [HH:MM:SS] の見出しを付ける
    （TimeCodeなしの切り方では False にして本文だけを並べる）。
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    _set_japanese_font(doc)

    doc.add_heading(title, level=0)
    if subtitle:
        sp = doc.add_paragraph(subtitle)
        sp_run = sp.runs[0]
        sp_run.italic = True
        sp_run.font.size = Pt(10)
        sp_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    for para in paragraphs:
        text = (para.get("text") or "").strip()
        if not text:
            continue
        if show_tc:
            tc_par = doc.add_paragraph()
            tc_run = tc_par.add_run(f"[{formats.hhmmss(para['start'])}]")
            tc_run.bold = True
            tc_run.font.size = Pt(9)
            tc_run.font.color.rgb = RGBColor(0x9C, 0x6B, 0x2B)
        body = doc.add_paragraph(text)
        body.paragraph_format.space_after = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------- Excel
def _safe_cell(text: str) -> str:
    """Excelの数式インジェクション対策。=,+,-,@ 等で始まる文字列は
    数式と解釈されないよう先頭に半角スペースを付ける。"""
    if text and text[0] in "=+-@\t\r":
        return " " + text
    return text


def _fill_sheet(ws, rows: list[dict]) -> None:
    from openpyxl.styles import Alignment, Font, PatternFill

    headers = ["No", "開始", "終了", "本文"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF", name="Yu Gothic")
        cell.fill = PatternFill("solid", fgColor="6B4A2B")
        cell.alignment = Alignment(horizontal="center", vertical="center")

    n = 0
    for r in rows:
        text = (r.get("text") or "").strip()
        if not text:
            continue
        n += 1
        ws.append([n, formats.hhmmss(r["start"]), formats.hhmmss(r["end"]), _safe_cell(text)])

    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 12
    ws.column_dimensions["D"].width = 95
    for row in ws.iter_rows(min_row=2):
        row[0].alignment = Alignment(horizontal="center", vertical="top")
        row[1].alignment = Alignment(horizontal="center", vertical="top")
        row[2].alignment = Alignment(horizontal="center", vertical="top")
        row[3].alignment = Alignment(wrap_text=True, vertical="top")
    ws.freeze_panes = "A2"


def build_xlsx(blocks: list[dict], segments: list[dict], body_label: str = "本文") -> bytes:
    """本文（選んだ切り方）と詳細（セグメント）の2シートを持つ Excel を bytes で返す。"""
    from openpyxl import Workbook

    wb = Workbook()
    ws1 = wb.active
    # シート名はExcelの制限（31文字以内・記号 : \ / ? * [ ] 不可）に収める
    ws1.title = body_label[:31]
    _fill_sheet(ws1, blocks)

    ws2 = wb.create_sheet("詳細（セグメント）")
    _fill_sheet(ws2, segments)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
